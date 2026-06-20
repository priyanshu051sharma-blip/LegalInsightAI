"""Document parsing and extraction"""
import os
from typing import Optional, Dict, Any, List
from pathlib import Path


class DocumentParser:
    """Parse legal documents"""

    @staticmethod
    def parse_pdf(file_path: str) -> Optional[str]:
        """Parse PDF file and extract text"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            return text
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return None

    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """Parse DOCX file and extract text"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return None

    @staticmethod
    def parse_txt(file_path: str) -> Optional[str]:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error parsing TXT: {e}")
            return None

    @staticmethod
    def parse_doc(file_path: str) -> Optional[str]:
        """Parse legacy DOC file"""
        try:
            import docx2txt
            return docx2txt.process(file_path)
        except Exception as e:
            print(f"Error parsing DOC: {e}")
            return None

    @staticmethod
    def parse_file(file_path: str, file_type: str) -> Optional[str]:
        """Parse file based on type"""
        if not os.path.exists(file_path):
            return None
        
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return DocumentParser.parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return DocumentParser.parse_docx(file_path)
        elif file_type == "txt":
            return DocumentParser.parse_txt(file_path)
        else:
            return None

    @staticmethod
    def extract_entities(text: str) -> Dict[str, List[str]]:
        """Extract legal entities from text"""
        try:
            import spacy
            
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text)
            
            entities = {
                "organizations": [],
                "persons": [],
                "dates": [],
                "locations": []
            }
            
            for ent in doc.ents:
                if ent.label_ == "ORG":
                    entities["organizations"].append(ent.text)
                elif ent.label_ == "PERSON":
                    entities["persons"].append(ent.text)
                elif ent.label_ == "DATE":
                    entities["dates"].append(ent.text)
                elif ent.label_ == "GPE":
                    entities["locations"].append(ent.text)
            
            # Remove duplicates
            for key in entities:
                entities[key] = list(set(entities[key]))
            
            return entities
        except Exception as e:
            print(f"Error extracting entities: {e}")
            return {}
